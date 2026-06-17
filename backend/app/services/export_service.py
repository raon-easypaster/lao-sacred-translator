import os
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportService:
    @staticmethod
    def to_markdown(data: dict) -> str:
        """Exports translation data to Markdown."""
        md = []
        md.append(f"# 번역 보고서: {data.get('direction', 'Lao Sacred Language Translation')}")
        md.append(f"**신뢰도(Confidence):** {data.get('confidence', 90)}%")
        md.append(f"**출처(Sources):** {data.get('source', 'N/A')}\n")
        md.append("---")
        
        md.append("## 1. 번역 결과")
        md.append(f"### [원문]\n{data.get('source_text', '')}\n")
        md.append(f"### [번역문]\n{data.get('translation', '')}\n")
        
        md.append("## 2. 스타일별 번역")
        md.append(f"- **직역(Literal):** {data.get('literal', 'N/A')}")
        md.append(f"- **의역(Contextual):** {data.get('contextual', 'N/A')}")
        md.append(f"- **설교체(Preaching):** {data.get('preaching', 'N/A')}\n")
        
        if data.get('cultural_warning'):
            md.append("## 3. 문화적 주의사항")
            md.append(f"> [!WARNING]\n> {data.get('cultural_warning')}\n")
            
        if data.get('commentary'):
            md.append("## 4. 용어 및 신학적 해설")
            md.append(f"{data.get('commentary')}\n")
            
        vocab = data.get('vocabulary', [])
        if vocab:
            md.append("## 5. 대조 용어 사전")
            md.append("| 한국어 | 일반 라오어 | 종교 라오어 | 왕실 라오어 | 의미 설명 |")
            md.append("|---|---|---|---|---|")
            for item in vocab:
                md.append(f"| {item.get('word','')} | {item.get('common','')} | {item.get('religious','')} | {item.get('royal','')} | {item.get('meaning','')} |")
                
        return "\n".join(md)

    @staticmethod
    def to_html(data: dict) -> str:
        """Exports translation data to HTML."""
        vocab_rows = ""
        vocab = data.get('vocabulary', [])
        for item in vocab:
            vocab_rows += f"""
            <tr>
                <td>{item.get('word','')}</td>
                <td class="lao">{item.get('common','')}</td>
                <td class="lao">{item.get('religious','')}</td>
                <td class="lao">{item.get('royal','')}</td>
                <td>{item.get('meaning','')}</td>
            </tr>
            """
            
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>LSLT 번역 보고서</title>
    <style>
        body {{ font-family: 'Inter', sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ color: #0b0e14; border-bottom: 2px solid #d4af37; padding-bottom: 10px; }}
        h2 {{ color: #161b26; margin-top: 30px; border-bottom: 1px solid #ddd; padding-bottom: 5px; }}
        .meta {{ background: #f8f9fa; padding: 15px; border-left: 4px solid #d4af37; margin-bottom: 20px; }}
        .block {{ background: #fdfdfd; border: 1px solid #eee; padding: 15px; border-radius: 4px; margin: 15px 0; }}
        .warning {{ background: #fff5f5; border-left: 4px solid #c92a2a; padding: 15px; color: #c92a2a; margin: 15px 0; }}
        .lao {{ font-family: 'Noto Sans Lao', sans-serif; font-size: 1.1em; color: #0066cc; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; }}
        th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
    <h1>Lao Sacred Language Translator 번역 보고서</h1>
    <div class="meta">
        <p><strong>번역 방향:</strong> {data.get('direction', 'Lao Sacred Language Translation')}</p>
        <p><strong>신뢰도(Confidence):</strong> {data.get('confidence', 90)}%</p>
        <p><strong>출처(Sources):</strong> {data.get('source', 'N/A')}</p>
    </div>
    
    <h2>1. 번역 결과</h2>
    <div class="block">
        <p><strong>[원문]</strong></p>
        <p>{data.get('source_text', '')}</p>
    </div>
    <div class="block">
        <p><strong>[번역문]</strong></p>
        <p class="lao">{data.get('translation', '')}</p>
    </div>
    
    <h2>2. 스타일별 번역</h2>
    <ul>
        <li><strong>직역:</strong> {data.get('literal', 'N/A')}</li>
        <li><strong>의역:</strong> {data.get('contextual', 'N/A')}</li>
        <li><strong>설교체:</strong> <span class="lao">{data.get('preaching', 'N/A')}</span></li>
    </ul>
    
    {"<h2>3. 문화적 주의사항</h2><div class='warning'>" + data.get('cultural_warning') + "</div>" if data.get('cultural_warning') else ""}
    
    {"<h2>4. 용어 및 신학적 해설</h2><div class='block'>" + data.get('commentary') + "</div>" if data.get('commentary') else ""}
    
    {"" if not vocab else f'''
    <h2>5. 대조 용어 사전</h2>
    <table>
        <thead>
            <tr>
                <th>한국어</th>
                <th>일반 라오어</th>
                <th>종교 라오어</th>
                <th>왕실 라오어</th>
                <th>의미 설명</th>
            </tr>
        </thead>
        <tbody>
            {vocab_rows}
        </tbody>
    </table>
    '''}
</body>
</html>
"""
        return html

    @staticmethod
    def to_docx(data: dict, output_path: str):
        """Exports translation data to a Word (.docx) file."""
        doc = DocxDocument()
        
        # Add title
        title = doc.add_paragraph()
        run = title.add_run("LSLT 번역 및 주해 보고서")
        run.bold = True
        run.font.size = Pt(20)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"번역 방향: {data.get('direction', 'N/A')}\n").bold = True
        meta.add_run(f"신뢰도: {data.get('confidence', 90)}%\n")
        meta.add_run(f"출처 문헌: {data.get('source', 'N/A')}")
        
        # Section 1
        doc.add_heading("1. 원문 및 번역문", level=1)
        p_src = doc.add_paragraph()
        p_src.add_run("[원문]\n").bold = True
        p_src.add_run(data.get('source_text', ''))
        
        p_trans = doc.add_paragraph()
        p_trans.add_run("[번역문]\n").bold = True
        p_trans.add_run(data.get('translation', ''))
        
        # Section 2
        doc.add_heading("2. 스타일별 대조 번역", level=1)
        doc.add_paragraph(f"직역: {data.get('literal', 'N/A')}")
        doc.add_paragraph(f"의역: {data.get('contextual', 'N/A')}")
        doc.add_paragraph(f"설교체: {data.get('preaching', 'N/A')}")
        
        # Section 3
        if data.get('cultural_warning'):
            doc.add_heading("3. 문화적 주의사항", level=1)
            p_warn = doc.add_paragraph()
            p_warn.add_run(data.get('cultural_warning'))
            
        # Section 4
        if data.get('commentary'):
            doc.add_heading("4. 신학적 및 개념적 해설", level=1)
            p_comm = doc.add_paragraph()
            p_comm.add_run(data.get('commentary'))
            
        # Section 5 (Glossary Table)
        vocab = data.get('vocabulary', [])
        if vocab:
            doc.add_heading("5. 매칭 전문 용어 사전", level=1)
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '한국어'
            hdr_cells[1].text = '일반 라오어'
            hdr_cells[2].text = '종교 라오어'
            hdr_cells[3].text = '왕실 라오어'
            hdr_cells[4].text = '설명'
            
            for item in vocab:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('word', '')
                row_cells[1].text = item.get('common', '')
                row_cells[2].text = item.get('religious', '')
                row_cells[3].text = item.get('royal', '')
                row_cells[4].text = item.get('meaning', '')
                
        doc.save(output_path)
